"""
Word document utilities for Word MCP Server.
Handles template rendering, placeholder scanning and find & replace.
"""

import json
import logging
import re
from io import BytesIO

from docx import Document
from docxtpl import DocxTemplate

logger = logging.getLogger(__name__)


def scan_placeholders(template_bytes: bytes) -> list[str]:
    """
    Scan a .docx template for Jinja2 {{ placeholder }} tags.

    Args:
        template_bytes: Raw bytes of the .docx file

    Returns:
        Sorted list of placeholder names found
    """
    doc = Document(BytesIO(template_bytes))
    placeholders = set()
    pattern = re.compile(r"\{\{\s*(\w+)\s*\}\}")

    # Scan paragraphs
    for para in doc.paragraphs:
        matches = pattern.findall(para.text)
        placeholders.update(matches)

    # Scan tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    matches = pattern.findall(para.text)
                    placeholders.update(matches)

    logger.info(f"Found {len(placeholders)} placeholders in template")
    return sorted(placeholders)


def render_template(template_bytes: bytes, context: dict) -> bytes:
    """
    Render a .docx template by filling in Jinja2 placeholders.

    Args:
        template_bytes: Raw bytes of the .docx template
        context: Dict of placeholder keys and their values

    Returns:
        Rendered .docx file as bytes
    """
    doc = DocxTemplate(BytesIO(template_bytes))
    doc.render(context)

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    logger.info(f"Template rendered with {len(context)} context values")
    return output.read()


def find_and_replace(docx_bytes: bytes, find_text: str, replace_text: str) -> tuple[bytes, int]:
    """
    Perform a find and replace on a rendered .docx file.

    Args:
        docx_bytes: Raw bytes of the .docx file
        find_text: Text to find
        replace_text: Text to replace it with

    Returns:
        Tuple of (updated .docx bytes, number of replacements made)
    """
    doc = Document(BytesIO(docx_bytes))
    replacements_made = 0

    # Search paragraphs
    for para in doc.paragraphs:
        if find_text in para.text:
            for run in para.runs:
                if find_text in run.text:
                    run.text = run.text.replace(find_text, replace_text)
                    replacements_made += 1

    # Search tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if find_text in para.text:
                        for run in para.runs:
                            if find_text in run.text:
                                run.text = run.text.replace(find_text, replace_text)
                                replacements_made += 1

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    logger.info(f"Find & replace: '{find_text}' → '{replace_text}' ({replacements_made} replacements)")
    return output.read(), replacements_made


def get_version_number(filename: str) -> int:
    """Extract version number from a filename like tender_BSCGlobal_v03.docx"""
    match = re.search(r'_v(\d+)\.docx$', filename, re.IGNORECASE)
    return int(match.group(1)) if match else 0


def get_family_name(filename: str) -> str:
    """Strip version suffix to get family name e.g. tender_BSCGlobal_v03.docx -> tender_BSCGlobal"""
    return re.sub(r'_v\d+\.docx$', '', filename, flags=re.IGNORECASE)


def next_version_filename(family_name: str, existing_files: list[str]) -> str:
    """
    Calculate the next version filename for a project.

    Args:
        family_name: The document project name e.g. tender_BSCGlobal
        existing_files: List of existing filenames in the output folder

    Returns:
        Next version filename e.g. tender_BSCGlobal_v03.docx
    """
    family_files = [
        f for f in existing_files
        if get_family_name(f).lower() == family_name.lower()
        and f.endswith(".docx")
    ]

    if not family_files:
        next_num = 1
    else:
        latest_version = max(get_version_number(f) for f in family_files)
        next_num = latest_version + 1

    return f"{family_name}_v{str(next_num).zfill(2)}.docx"


def get_latest_version_filename(family_name: str, existing_files: list[str]) -> str | None:
    """
    Get the latest version filename for a project.

    Args:
        family_name: The document project name e.g. tender_BSCGlobal
        existing_files: List of existing filenames in the output folder

    Returns:
        Latest version filename or None if no versions exist
    """
    family_files = [
        f for f in existing_files
        if get_family_name(f).lower() == family_name.lower()
        and f.endswith(".docx")
    ]

    if not family_files:
        return None

    return max(family_files, key=get_version_number)


def serialize_context(context: dict) -> str:
    """Serialize a context dict to a JSON string for storage."""
    return json.dumps(context, indent=2, ensure_ascii=False)


def deserialize_context(json_str: str) -> dict:
    """Deserialize a JSON string back to a context dict."""
    if not json_str:
        return {}
    return json.loads(json_str)